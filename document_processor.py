import os
import logging
import json
from datetime import datetime
import sqlite3
import sys
import textract
from pypdf import PdfReader
from docx import Document as DocxDocument
from sqlalchemy import create_engine
from sqlalchemy.orm import sessionmaker

# Configuration du logging
logging.basicConfig(level=logging.INFO)
logger = logging.getLogger(__name__)

def extract_text_from_pdf(filepath):
    """
    Extrait le texte d'un fichier PDF
    """
    try:
        if not os.path.exists(filepath):
            logger.error(f"Le fichier {filepath} n'existe pas")
            return ""
        
        reader = PdfReader(filepath)
        text = ""
        for page in reader.pages:
            page_text = page.extract_text()
            if page_text:
                text += page_text + "\n\n"
        
        return text
    except Exception as e:
        logger.error(f"Erreur lors de l'extraction du texte du PDF {filepath}: {str(e)}")
        return ""

def extract_text_from_docx(filepath):
    """
    Extrait le texte d'un fichier Word DOCX
    """
    try:
        if not os.path.exists(filepath):
            logger.error(f"Le fichier {filepath} n'existe pas")
            return ""
        
        doc = DocxDocument(filepath)
        text = ""
        for para in doc.paragraphs:
            if para.text:
                text += para.text + "\n"
        
        # Extraire le texte des tableaux
        for table in doc.tables:
            for row in table.rows:
                row_text = []
                for cell in row.cells:
                    if cell.text:
                        row_text.append(cell.text.strip())
                if row_text:
                    text += " | ".join(row_text) + "\n"
        
        return text
    except Exception as e:
        logger.error(f"Erreur lors de l'extraction du texte du DOCX {filepath}: {str(e)}")
        return ""

def extract_text_using_textract(filepath):
    """
    Utilise textract pour extraire le texte de divers formats de fichiers
    """
    try:
        if not os.path.exists(filepath):
            logger.error(f"Le fichier {filepath} n'existe pas")
            return ""
        
        # Textract prend en charge de nombreux formats: DOC, DOCX, XLS, XLSX, CSV, TXT, RTF, etc.
        text = textract.process(filepath)
        
        # Convertir les bytes en string
        if isinstance(text, bytes):
            text = text.decode('utf-8', errors='replace')
        
        return text
    except Exception as e:
        logger.error(f"Erreur lors de l'extraction du texte via textract pour {filepath}: {str(e)}")
        return ""

def process_document(document_id):
    """
    Traite un document pour en extraire le contenu et le stocke dans la base de données
    """
    try:
        # Connexion à la base de données
        from main import app
        import os
        
        # Importer les models seulement ici pour éviter l'import circulaire
        from models import Document
        
        # Récupérer le document depuis la base de données
        with app.app_context():
            document = Document.query.get(document_id)
            if not document:
                logger.error(f"Document avec l'ID {document_id} non trouvé")
                return None
        
        # Construire le chemin complet vers le fichier
        filepath = os.path.join("static/uploads", document.filepath)
        
        # Extraire le texte selon le type de fichier
        text = ""
        file_lower = filepath.lower()
        
        try:
            # Traitement spécifique pour certains formats de fichiers
            if file_lower.endswith(".pdf"):
                text = extract_text_from_pdf(filepath)
            elif file_lower.endswith(".docx"):
                text = extract_text_from_docx(filepath)
            elif file_lower.endswith((".txt", ".csv")):
                # Lire directement les fichiers texte
                with open(filepath, 'r', encoding='utf-8', errors='replace') as f:
                    text = f.read()
            else:
                # Utiliser textract pour tous les autres formats (xls, xlsx, doc, rtf, etc.)
                logger.info(f"Extraction du texte via textract pour {filepath}")
                text = extract_text_using_textract(filepath)
        except Exception as e:
            logger.error(f"Erreur lors de l'extraction du texte de {filepath}: {str(e)}")
            
        # Si le texte est vide, essayer avec textract en dernier recours
        if not text:
            try:
                logger.info(f"Tentative d'extraction de secours avec textract pour {filepath}")
                text = extract_text_using_textract(filepath)
            except Exception as e:
                logger.error(f"Échec de l'extraction de secours pour {filepath}: {str(e)}")
                
        # Si aucun texte n'a été extrait, retourner None
        if not text:
            logger.warning(f"Aucun texte extrait du fichier {filepath}")
            return None
        
        # Mettre à jour le document avec le contenu extrait
        if text:
            # Stocker le contenu du document dans un fichier JSON
            content_dir = "static/document_contents"
            os.makedirs(content_dir, exist_ok=True)
            
            content_filename = f"{document.id}_{datetime.now().strftime('%Y%m%d%H%M%S')}.json"
            content_path = os.path.join(content_dir, content_filename)
            
            # Préparer les données à stocker dans le JSON
            document_data = {
                "document_id": document.id,
                "filename": document.filename,
                "content": text,
                "extracted_at": datetime.now().isoformat(),
                "document_type": document.document_type,
                "document_category": document.document_category,
                "document_date": document.document_date.isoformat() if document.document_date else None,
                "amount": document.amount,
                "description": document.description
            }
            
            # Ajouter les identifiants property_id et company_id s'ils existent
            if document.property_id:
                document_data["property_id"] = document.property_id
            
            if hasattr(document, 'company_id') and document.company_id:
                document_data["company_id"] = document.company_id
            
            # Sauvegarder le document JSON
            with open(content_path, 'w', encoding='utf-8') as f:
                json.dump(document_data, f, ensure_ascii=False, indent=4)
            
            logger.info(f"Contenu extrait et enregistré pour le document {document.id}: {document.filename}")
            return content_path
        else:
            logger.warning(f"Aucun contenu extrait du document {document.id}: {document.filename}")
            return None
    
    except Exception as e:
        logger.error(f"Erreur lors du traitement du document {document_id}: {str(e)}")
        return None

def process_all_documents():
    """
    Traite tous les documents de la base de données
    """
    # Importation des modules nécessaires
    from main import app
    from models import Document
    
    # Utilisation du contexte d'application Flask
    with app.app_context():
        documents = Document.query.all()
        results = {
            "success": [],
            "failed": []
        }
    
    # Liste des extensions de fichiers supportées
    supported_extensions = [
        ".pdf", ".docx", ".doc", ".xls", ".xlsx", ".txt", ".csv", 
        ".rtf", ".odt", ".ppt", ".pptx", ".html", ".htm"
    ]
    
    for document in documents:
        try:
            filepath = os.path.join("static/uploads", document.filepath)
            
            # Vérifier si l'extension du fichier est supportée
            file_extension = os.path.splitext(filepath.lower())[1]
            
            if file_extension in supported_extensions or len(file_extension) == 0:
                logger.info(f"Traitement du document {document.id}: {document.filename}")
                content_path = process_document(document.id)
                
                if content_path:
                    results["success"].append({
                        "id": document.id,
                        "filename": document.filename,
                        "content_path": content_path
                    })
                else:
                    results["failed"].append({
                        "id": document.id,
                        "filename": document.filename,
                        "reason": "Échec de l'extraction du contenu"
                    })
            else:
                logger.warning(f"Type de fichier non supporté pour {document.filename} (extension: {file_extension})")
                results["failed"].append({
                    "id": document.id,
                    "filename": document.filename,
                    "reason": f"Type de fichier non supporté (extension: {file_extension})"
                })
        except Exception as e:
            logger.error(f"Erreur lors du traitement du document {document.id}: {str(e)}")
            results["failed"].append({
                "id": document.id,
                "filename": document.filename,
                "reason": str(e)
            })
    
    logger.info(f"Traitement terminé. {len(results['success'])} documents traités avec succès, {len(results['failed'])} échecs.")
    return results

def get_document_content(document_id):
    """
    Récupère le contenu d'un document depuis son fichier JSON
    """
    try:
        content_dir = "static/document_contents"
        if not os.path.exists(content_dir):
            return None
        
        # Chercher tous les fichiers commençant par l'ID du document
        prefix = f"{document_id}_"
        content_files = [f for f in os.listdir(content_dir) if f.startswith(prefix) and f.endswith('.json')]
        
        if not content_files:
            # Si aucun fichier n'existe déjà, tenter de traiter le document
            process_document(document_id)
            # Chercher à nouveau
            content_files = [f for f in os.listdir(content_dir) if f.startswith(prefix) and f.endswith('.json')]
            
        if content_files:
            # Prendre le fichier le plus récent (tri alphabétique, le timestamp fait partie du nom)
            content_files.sort(reverse=True)
            latest_file = content_files[0]
            
            with open(os.path.join(content_dir, latest_file), 'r', encoding='utf-8') as f:
                return json.load(f)
        
        return None
    
    except Exception as e:
        logger.error(f"Erreur lors de la récupération du contenu du document {document_id}: {str(e)}")
        return None

def get_all_documents_content():
    """
    Récupère le contenu de tous les documents pour les fournir à l'assistant IA
    """
    try:
        content_dir = "static/document_contents"
        if not os.path.exists(content_dir):
            os.makedirs(content_dir, exist_ok=True)
            return []
        
        documents_content = []
        
        for filename in os.listdir(content_dir):
            if filename.endswith('.json'):
                try:
                    with open(os.path.join(content_dir, filename), 'r', encoding='utf-8') as f:
                        content = json.load(f)
                        documents_content.append(content)
                except Exception as e:
                    logger.error(f"Erreur lors de la lecture du fichier {filename}: {str(e)}")
        
        return documents_content
    
    except Exception as e:
        logger.error(f"Erreur lors de la récupération du contenu de tous les documents: {str(e)}")
        return []

def search_in_documents(query, property_id=None, company_id=None):
    """
    Recherche une requête dans le contenu des documents
    """
    try:
        # Convertir la requête en minuscules pour une recherche insensible à la casse
        query = query.lower()
        
        # Récupérer tous les contenus de documents
        all_contents = get_all_documents_content()
        
        # Filtrer par propriété si demandé
        if property_id:
            all_contents = [content for content in all_contents if content.get("property_id") == property_id]
        
        # Filtrer par société si demandé
        if company_id:
            all_contents = [content for content in all_contents if content.get("company_id") == company_id]
        
        # Rechercher la requête dans le contenu
        results = []
        for content in all_contents:
            document_text = content.get("content", "").lower()
            if query in document_text:
                # Trouver les extraits contenant la requête
                excerpts = []
                lines = document_text.split('\n')
                for line in lines:
                    if query in line:
                        excerpts.append(line.strip())
                
                if excerpts:
                    result = {
                        "document_id": content.get("document_id"),
                        "filename": content.get("filename"),
                        "excerpts": excerpts[:5],  # Limiter à 5 extraits par document
                        "document_type": content.get("document_type"),
                        "document_category": content.get("document_category"),
                        "document_date": content.get("document_date"),
                        "amount": content.get("amount"),
                        "description": content.get("description")
                    }
                    
                    # Ajouter property_id et company_id s'ils existent
                    if content.get("property_id"):
                        result["property_id"] = content.get("property_id")
                    
                    if content.get("company_id"):
                        result["company_id"] = content.get("company_id")
                    
                    results.append(result)
        
        return results
    
    except Exception as e:
        logger.error(f"Erreur lors de la recherche dans les documents: {str(e)}")
        return []

def get_property_documents_content(property_id):
    """
    Récupère le contenu de tous les documents associés à une propriété
    """
    try:
        # Récupérer tous les contenus de documents
        all_contents = get_all_documents_content()
        
        # Filtrer par propriété
        property_contents = [content for content in all_contents if content.get("property_id") == property_id]
        
        return property_contents
    
    except Exception as e:
        logger.error(f"Erreur lors de la récupération du contenu des documents de la propriété {property_id}: {str(e)}")
        return []

def get_company_documents_content(company_id):
    """
    Récupère le contenu de tous les documents associés à une société
    """
    try:
        # Récupérer tous les contenus de documents
        all_contents = get_all_documents_content()
        
        # Filtrer par société
        company_contents = [content for content in all_contents if content.get("company_id") == company_id]
        
        return company_contents
    
    except Exception as e:
        logger.error(f"Erreur lors de la récupération du contenu des documents de la société {company_id}: {str(e)}")
        return []